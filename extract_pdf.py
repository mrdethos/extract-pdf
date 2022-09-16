import os
import fitz
import io
import argparse
import pytesseract
from docx import Document
from docx.shared import Inches, Pt
from PIL import Image

pytesseract.pytesseract.tesseract_cmd = r'C:\Users\daniel.campos\AppData\Local\Tesseract-OCR\tesseract.exe'

''' Extrai textos gerais e textos de imagens de um PDF '''
def get_texts(args):
    # verifica se foi inserido um arquivo .pdf ou um diretório
    if ".pdf" in args.path:
        pdf_file = fitz.open(args.path)
        string = ''
        # percorre todas as páginas do pdf
        for i in range(len(pdf_file)):
            page = pdf_file[i]
            # pega todos os textos da página
            string += page.get_text()
            # pega todas as imagens da página
            image_list = page.get_images()
            for image_index, img in enumerate(page.get_images(), start=1):
                # xref da imagem
                xref = img[0]
                # bytes da imagem
                base_image = pdf_file.extract_image(xref)
                image_bytes = base_image["image"]
                # carrega imagem no PIL e usa a biblioteca tesseract para OCR
                with Image.open(io.BytesIO(image_bytes)) as image:
                    string += pytesseract.image_to_string(image, lang = 'por')
        # deixa apenas o nome do arquivo na string
        if '/' in args.path:
            name = args.path
            count = name.count('/')
            for i in range(count):
                before, sep, name = name.partition('/')
        elif '\\' in args.path:
            name = args.path
            count = name.count('\\')
            for i in range(count):
                before, sep, name = name.partition('\\')
        save_texts(string, name)
    else:
        for (root, dirs, files) in os.walk(args.path, topdown=True):
            for name in files:
                pdf_file = fitz.open(args.path+name)
                string = ''
                # percorre todas as páginas do pdf
                for i in range(len(pdf_file)):
                    page = pdf_file[i]
                    # pega todos os textos da página
                    string += page.get_text()
                    # pega todas as imagens da página
                    image_list = page.get_images()
                    for image_index, img in enumerate(page.get_images(), start=1):
                        # xref da imagem
                        xref = img[0]
                        # bytes da imagem
                        base_image = pdf_file.extract_image(xref)
                        image_bytes = base_image["image"]
                        # carrega imagem no PIL e usa a biblioteca tesseract para OCR
                        with Image.open(io.BytesIO(image_bytes)) as image:
                            string += pytesseract.image_to_string(image, lang = 'por')
                save_texts(string, name)


''' Salva os textos em um arquivo .docx '''
def save_texts(string, filename):
    doc = Document()
    # adicionando estilização ao documento
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.18)
        section.bottom_margin = Inches(0.78)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.78)
    body_style = doc.styles['Body Text']
    body = doc.add_paragraph(style=body_style).add_run(f'{string}')
    body.font.size = Pt(12)
    body.font.name = 'Arial'
    filename = filename.replace(".pdf", ".docx")
    print(filename, " salvo com sucesso.")
    doc.save('results/'+filename)

def parse_args():
    parser = argparse.ArgumentParser()
    parser.add_argument('path', type=str,
                        help='Caminho do .pdf ex:"\pdf\imagem.pdf"')
    parser.add_argument('-a', '--all', action='store_true',
                        help='Extrai todos os arquivos de um caminho')
    return parser.parse_args()

if __name__ == '__main__':
    args = parse_args()
    get_texts(args)
